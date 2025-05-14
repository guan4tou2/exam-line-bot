import re,random,json,os
from docx import Document
from flask import Flask, request
from dotenv import load_dotenv

# 載入 LINE Message API 相關函式庫
from linebot import LineBotApi, WebhookHandler
from linebot.exceptions import InvalidSignatureError
from linebot.models import MessageEvent, TextMessage, TextSendMessage,PostbackAction,URIAction, MessageAction, TemplateSendMessage, ButtonsTemplate,FlexSendMessage

app = Flask(__name__)

access_token = load_dotenv(os.getenv('ACCESS_TOKEN'))
secret = load_dotenv(os.getenv('SECRET'))
char_list = list("ABCD")

def get_question(file):
    document=Document(file)
    # print("題目數量:",randommax:=len(document.paragraphs))
    randommax=len(document.paragraphs)
    random_number=random.randrange(0,randommax-1)
    question_pattern = re.compile(r"\(([A-D])\)\d{1,2}\.(.*？)\(A\)(.*)\(B\)(.*)\(C\)(.*)\(D\)(.*)。", re.S)

    match = question_pattern.findall(document.paragraphs[random_number].text)
    if match:
        question_number, question_text, option1, option2, option3, option4 = match[0]
        question_data = {
            "question_answer": question_number,
            "question_text": question_text.strip().replace('\n',''),
            "options": {
                "A": option1.strip().replace('\n',''),
                "B": option2.strip().replace('\n',''),
                "C": option3.strip().replace('\n',''),
                "D": option4.strip().replace('\n','')
            }
        }
    print(random_number)
    return question_data

def topic(file):
    question_data=get_question(file)   
    question_text = question_data['question_text']
    question_answer = question_data['question_answer']
    question_options = question_data['options']
    # random.shuffle(char_list)       
    flex_message=json.load(open('flex_message.json','r',encoding='utf-8'))
    flex_message["body"]["contents"][0]["text"]=question_text
    flex_message["footer"]["contents"][0]["contents"][0]["text"]=question_options[char_list[0]]
    flex_message["footer"]["contents"][1]["contents"][0]["text"]=question_options[char_list[1]]
    flex_message["footer"]["contents"][2]["contents"][0]["text"]=question_options[char_list[2]]
    flex_message["footer"]["contents"][3]["contents"][0]["text"]=question_options[char_list[3]]

    # print(question_answer)
    flex_message["footer"]["contents"][ord(question_answer)-ord("A")]['backgroundColor']="#B3A398"
    # print(char_list)
    return FlexSendMessage('iPAS題目',flex_message)


@app.route("/", methods=['POST'])
def linebot():
    body = request.get_data(as_text=True)    
    # question_data=get_question("資訊安全管理概論.docx")   
    # question_text = question_data['question_text']
    # question_answer = question_data['question_answer']
    # question_options = question_data['options']
    # # random.shuffle(char_list)       
    # flex_message=json.load(open('flex_message.json','r',encoding='utf-8'))
    
    try:
        json_data = json.loads(body)                         # json 格式化訊息內容
        # access_token = access_token
        # secret = secret
        line_bot_api = LineBotApi(access_token)              # 確認 token 是否正確
        handler = WebhookHandler(secret)                     # 確認 secret 是否正確
        signature = request.headers['X-Line-Signature']      # 加入回傳的 headers
        handler.handle(body, signature)                      # 綁定訊息回傳的相關資訊
        tk = json_data['events'][0]['replyToken']            # 取得回傳訊息的 Token
        type = json_data['events'][0]['message']['type']     # 取得 LINe 收到的訊息類型
        print(f"type:{type}")
        if type=='text':
            msg = json_data['events'][0]['message']['text']  # 取得 LINE 收到的文字訊息
            print(f"msg:{msg}")                                       # 印出內容
            reply = msg
        else:
            reply = '你傳的不是文字呦～'
        

        # flex_message["body"]["contents"][0]["text"]=question_text
        # flex_message["footer"]["contents"][0]["contents"][0]["text"]=question_options[char_list[0]]
        # flex_message["footer"]["contents"][1]["contents"][0]["text"]=question_options[char_list[1]]
        # flex_message["footer"]["contents"][2]["contents"][0]["text"]=question_options[char_list[2]]
        # flex_message["footer"]["contents"][3]["contents"][0]["text"]=question_options[char_list[3]]

        # flex_message["footer"]["contents"][ord(question_answer)-ord("A")]['backgroundColor']="#B3A398"
 
        reply_arr=[]
        # reply_arr.append(TextSendMessage(f"題目\n{question_text}"))
        # reply_arr.append(TextSendMessage(f"A：{question_options[char_list[0]]}\n\nB：{question_options[char_list[1]]}\n\nC：{question_options[char_list[2]]}\n\nD：{question_options[char_list[3]]}"))
        # reply_arr.append(TextSendMessage(f"{question_options[question_answer]}"))
        # reply_arr.append(FlexSendMessage('topic',flex_message))
        reply_arr.append(topic("資訊安全管理概論.docx"))
        line_bot_api.reply_message(tk,reply_arr)# 回傳訊息


    except:
        print(body)                                          # 如果發生錯誤，印出收到的內容
    return 'OK'                                              # 驗證 Webhook 使用，不能省略

if __name__ == "__main__":
    app.run()
